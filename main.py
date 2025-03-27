# streamlit run "C:/Users/keno/OneDrive/Documents/Projects/n_secondbrain/1_TESTAGAIN.py"
# Create me a detailed guide and plan for me as an Indonesian to be able to work or be a university student (Master's) in Japan.

import streamlit as st
import requests
from bs4 import BeautifulSoup, ParserRejectedMarkup
from newspaper import Article
import google.generativeai as genai
import pandas as pd
import datetime
import time
import threading
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import matplotlib.pyplot as plt
from io import BytesIO
import zipfile
from docx import Document
from PIL import Image
import base64
import random
from urllib.parse import quote_plus
import nltk
nltk.download('punkt_tab')

# ============================================== HTML AND CSS =========================================================================================

logo = Image.open("logo.png")

st.set_page_config(
    page_title="n_secondbrain",
    page_icon=logo,
    layout="wide",  # or "wide" if you prefer
    initial_sidebar_state="auto"
)

st.set_option('client.showErrorDetails', False)

st.markdown(
    """
    <style>
    section[data-testid="stMain"] > div[data-testid="stMainBlockContainer"] {
         padding-top: 0px;  # Remove padding completely
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Hide Streamlit style elements
hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                header {visibility: hidden;}
                </style>
                """
st.markdown(hide_st_style, unsafe_allow_html=True)

st.markdown("""
    <style>
    [data-testid="stTextArea"] {
        color: #C8ACD6;
    }
    </style>
    """, unsafe_allow_html=True)

# Set Montserrat font
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Montserrat', sans-serif;
    }
    </style>
    """, unsafe_allow_html=True)

# Change color of specific Streamlit elements
st.markdown("""
    <style>
    .st-emotion-cache-1o6s5t7 {
        color: #ababab !important;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stExpander {
        background-color: #C8ACD6;
        border-radius: 10px;
    }
    
    .stExpander > details {
        background-color: #C8ACD6;
        border-radius: 10px;
    }
    
    .stExpander > details > summary {
        background-color: #C8ACD6;
        border-radius: 10px 10px 0 0;
        padding: 10px;
    }
    
    .stExpander > details > div {
        background-color: #C8ACD6;
        border-radius: 0 0 10px 10px;
        padding: 10px;
    }
    
    .stCheckbox {
        background-color: #C8ACD6;
        border-radius: 5px;
        padding: 5px;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stButton > button {
        color: #FFFFFF;
        background-color: #2E236C;
        border-radius: 10px;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown(
    """
    <style>
    .streamlit-expanderHeader {
        font-size: 20px;
    }
    .streamlit-expanderContent {
        max-height: 400px;
        overflow-y: auto;
    }
    </style>
    """,
    unsafe_allow_html=True
)

def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

def set_png_as_page_bg(png_file):
    bin_str = get_base64_of_bin_file(png_file)
    page_bg_img = '''
    <style>
    .stApp {
        background-image: url("data:image/png;base64,%s");
        background-size: cover;
    }
    </style>
    ''' % bin_str
    
    st.markdown(page_bg_img, unsafe_allow_html=True)

# Set the background image
set_png_as_page_bg("background.jpg")

# ============================================== FUNCTIONS =========================================================================================

# --- Second-Brain LLM Functions ---
def configure_model():
    if "model" not in st.session_state:
        apisheetskey = "1sIEI-_9N96ndRJgWDyl0iL65bACeGQ74MncOV4HQCXY"
        url_apikey = f'https://docs.google.com/spreadsheet/ccc?key={apisheetskey}&output=csv'
        df_apikey = pd.read_csv(url_apikey)
        platform = "Gemini"
        email = "kenoelle11@aiesec.net"
        apikeyxloc = df_apikey['Platform'].str.contains(platform).idxmax()
        apikeyxloc = df_apikey['Platform'].str.contains(email).idxmax()
        apikey = df_apikey.iloc[apikeyxloc, 2]
        st.session_state["apikey"] = apikey
        genai.configure(api_key=apikey)
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 60,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }
        st.session_state["model"] = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config=generation_config
        )
        st.session_state["generation_config"] = generation_config
    else:
        st.error("Model did not load successfully, please reset the website.")
    
    return st.session_state.get("model")

# ---------------------------------------
# Utility Functions
# ---------------------------------------

def get_random_headers():
    """Return a random headers dictionary for web requests."""
    user_agents = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/101.0.4951.54 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 Version/15.1 Safari/605.1.15",
        "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 Chrome/95.0.4638.69 Safari/537.36",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 Chrome/100.0.4896.127 Safari/537.36"
    ]
    headers = {
        "User-Agent": random.choice(user_agents),
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Referer": "https://www.google.com/"
    }
    return headers

# ---------------------------------------
# News & Article Extraction Functions
# ---------------------------------------

def getNewsData(query):
    """Searches for news articles based on the given query, with random headers."""
    headers = {
        "User-Agent": random.choice([
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/101.0.4951.54 Safari/537.36",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 Version/15.1 Safari/605.1.15",
            "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 Chrome/95.0.4638.69 Safari/537.36",
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 Chrome/100.0.4896.127 Safari/537.36"
        ]),
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Referer": "https://www.google.com/"
    }

    encoded_query = quote_plus(query)
    url = f"https://www.google.com/search?q={encoded_query}&gl=us&tbm=nws&num=15"

    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()  # Raise an exception for bad status codes

        try:
            soup = BeautifulSoup(response.content, "html.parser")
        except ParserRejectedMarkup as e:
            print(f"Error parsing HTML: {e}")
            return []

        news_results = []

        for el in soup.select("div.SoaBEf"):
            news_results.append(
                {
                    "link": el.find("a")["href"],
                    "title": el.select_one("div.MBeuO").get_text(),
                    "snippet": el.select_one(".GI74Re").get_text(),
                    "date": el.select_one(".LfVVr").get_text(),
                    "source": el.select_one(".NUnG9d span").get_text(),
                }
            )
        return news_results

    except requests.exceptions.RequestException as e:
        print(f"Error fetching news data: {e}")
        return []

def download_and_parse_with_timeout(article, link, timeout=10):
    """Downloads and parses an article with a time limit, filters short content."""
    
    def download_and_parse():
        try:
            article.download()
            article.parse()
            article.nlp()  # Extract summary, keywords, etc.
            if len(article.text) >= 1000: # Check content length
                result['data'] = {
                    "title": article.title,
                    "link": link,
                    "summary": article.summary,
                    "content": article.text
                }
            else:
                result['data'] = None # return none if content too short
        except Exception as e:
            result['error'] = str(e)

    result = {'data': None, 'error': None}
    thread = threading.Thread(target=download_and_parse)
    thread.start()
    thread.join(timeout=timeout)

    if thread.is_alive():
        print(f"Download and parse timed out after {timeout} seconds: {link}")
        return None  # Return None to indicate timeout

    if result['error']:
        print(f"Error downloading or parsing {link}: {result['error']}")
        return None

    return result['data']

def extract_article_info(links):
    """Extracts relevant information from a list of article links, filters short articles."""
    articles_info = []
    for link in links:
        try:
            article = Article(link)
            article_data = download_and_parse_with_timeout(article, link)
            if article_data: # only append if article data is not none (not short)
                articles_info.append(article_data)

        except Exception as e:
            print(f"Sorry, article not downloadable: {link}")
            print(f"Error: {e}")

    return articles_info

# ---------------------------------------
# Query Generation Functions
# ---------------------------------------

def make_prompt_queries_for_part(prompt, n_queries=5, model=None):
    """
    Uses the language model to generate n_queries short (1-2 word) Google dork queries for a single TOC part.
    The prompt instructs the model to output one query per line.
    """
    full_prompt = f"""
    For this prompt: {prompt}
    Generate {n_queries} short Google dorking search queries. 
    The queries should only be for CODING related subjects.
    Each query should be 1-2 words only, without quotation marks except for specific names.
    Format each query on a new line.
    The queries should exclude filetypes : xlsx,pdf,docx,ppt, etc.
    """
    response = model.generate_content(full_prompt)
    queries = [line.strip() for line in response.text.strip().split('\n') if line.strip()]
    return queries[:n_queries]

def make_prompt_queries(prompt, n_parts, model):
    """
    Generates Google search queries for each part of the Table of Contents (TOC).
    Handles potential variations in LLM response format to prevent index errors.
    """
    queries_by_part = {}
    for i in range(n_parts):
        part_prompt = f"""
            For this main prompt: {prompt}, generate 5 google dork queries. 
            These queries must be for part {i + 1} of the prompt.
            The queries should only be for CODING related subjects.
            Search/dork should be about the key words of the project of the prompt.
            The queries should exclude filetypes : xlsx,pdf,docx,ppt, etc.
            These queries must be very specific, and should not use quotation marks unless it is for a name or institution.
            Do not mention the other parts.
            Do not include any other text except for the queries.
            Follow this format:
            1. Query 1
            2. Query 2
            3. Query 3
            4. Query 4
            5. Query 5
        """
        try:
            response = model.generate_content(part_prompt)
            lines = response.text.strip().split('\n')
            queries = []
            for line in lines:
                line = line.strip()
                if line:  # Check if line is not empty
                    parts = line.split(". ", 1)
                    if len(parts) > 1:
                        queries.append(parts[1].strip())
                    else:
                        # Handle cases where the line doesn't match the expected format
                        queries.append(line)  # Use the whole line as a query
            queries_by_part[f"Part {i + 1}"] = queries[:5] #take only the first 5 queries
        except Exception as e:
            st.error(f"Error generating queries for Part {i + 1}: {e}")
            queries_by_part[f"Part {i+1}"] = ["Error: Could not generate Queries"]
    return queries_by_part

# ---------------------------------------
# TOC & Prompt Generation Functions
# ---------------------------------------

def table_of_contents(prompt, n_parts=5, model=None):
    if "model" not in st.session_state and model is None:
        configure_model()
        model = st.session_state["model"]

    if model is None:
        st.error("Model did not load successfully, please reset the website.")
        return ""

    full_prompt = f"""
        For this prompt: {prompt}
        The language of the code is {language}"
        Divide the task in the prompt into {n_parts} parts.
        Remember, this is a CODING task, the order of the sequential process matters and will affect, choose wisely.
        Generate a short title for each part.
        Just generate the list of titles, nothing else.
        Follow the format here : 
        Part 1: (title)
        Part 2: (title)
        Make sure every part is separated by a spaced line, and every line can only have 1 part.
    """

    response = model.generate_content(full_prompt)
    toc_text = response.text.strip()
    markdown_output = toc_text.replace('\n', '  \n')  # Add double space for markdown line breaks
    return markdown_output

def extract_parts(toc_text):
    """
    Memisahkan string TOC menjadi daftar bagian berdasarkan baris baru.
    Tidak memvalidasi format konten setiap baris.
    """
    parts = toc_text.strip().split('\n')
    parts = [part.strip() for part in parts if part.strip()] #Remove leading and trailing spaces, and empty lines.
    return parts


def generate_prompts_for_parts(parts_list, model, global_context):
    """
    Generates detailed prompts for each part of the Table of Contents (TOC).

    Args:
        parts_list (list): List of TOC part names.
        model (genai.GenerativeModel): The generative AI model.
        global_context (str): The overall context for the task.

    Returns:
        list: A list of detailed prompts, one for each TOC part.
    """
    prompts = []
    for part in parts_list:
        full_prompt = f"""
        {global_context}
        The language of the code is {language}"
        For this part of the task: "{part}", please generate a detailed prompt. 
        The detailed prompt should be very specific, and should include what the user should focus on for this part.
        Do not mention the other parts.
        """
        try:
            response = model.generate_content(full_prompt)
            prompts.append(response.text.strip())
        except Exception as e:
            st.error(f"Error generating prompt for part '{part}': {e}")
            prompts.append(f"Error: Could not generate detailed prompt for {part}") #add error message to prompt list
    return prompts

def llm_generate(prompt, model=None, context=None):
    """
    Generates content using the language model.
    Optionally includes additional context.
    """
    if context:
        prompt_with_context = f"{prompt}\n\nContext:\n{context}"
        response = model.generate_content(prompt_with_context)
    else:
        response = model.generate_content(prompt)
    text = response.text.replace("*", "")
    return text

def extract_key_points(text_list):
    """
    Concatenates a list of text strings into a single string with line breaks.
    """
    return "\n".join(text_list)

# ---------------------------------------
# Embedding and Visualization Functions
# ---------------------------------------

def embed(text, existing_vectorizer=None):
    """
    Computes TF-IDF embeddings for the given text.
    If an existing vectorizer is provided, uses it to transform the text.
    """
    if existing_vectorizer is None:
        vectorizer = TfidfVectorizer()
        embeddings = vectorizer.fit_transform([text]).toarray()
        return embeddings, vectorizer
    else:
        embeddings = existing_vectorizer.transform([text]).toarray()
        return embeddings

def visualize_iterative_embeddings(embeddings):
    """
    Visualizes the cosine similarity between consecutive iterations.
    """
    num_iterations = len(embeddings)
    similarities = []
    for i in range(num_iterations - 1):
        similarity = cosine_similarity(embeddings[i], embeddings[i + 1])[0][0]
        similarities.append(similarity)
    plt.figure(figsize=(10, 6))
    plt.fill_between(range(1, num_iterations), similarities, color='skyblue', alpha=0.7)
    plt.plot(range(1, num_iterations), similarities, marker='o', linestyle='-', color='blue')
    plt.xlabel("Iteration")
    plt.ylabel("Cosine Similarity (between consecutive iterations)")
    plt.title("Iterative Refinement: Embedding Similarity over Iterations")
    plt.grid(True)
    plt.show()

# ---------------------------------------
# Iterative Refinement Function
# ---------------------------------------

def iterative_refinement(initial_prompt, internet_knowledge, previous_outputs=[], iterations=5, global_context="", model=None):
    """
    Refines the initial prompt iteratively, using different 15,000-character chunks of internet_knowledge.
    Includes previous outputs as context.
    Returns the final refined output, the initial output, and all thinking logs.
    """
    knowledge_len = len(internet_knowledge)
    chunk_size = 15000
    
    thinking_logs = []
    all_responses = []
    
    # Handle the first iteration separately to get the initial response
    start_index = 0
    end_index = min(start_index + chunk_size, knowledge_len)
    truncated_knowledge = internet_knowledge[start_index:end_index]

    previous_output_context = "\n\n".join(previous_outputs) if previous_outputs else ""
    combined_prompt = f"{global_context}\n\n{truncated_knowledge}\n\n{previous_output_context}\n\n{initial_prompt}" if global_context or truncated_knowledge or previous_output_context else initial_prompt

    initial_response = llm_generate(combined_prompt, model=model)
    all_responses.append(initial_response)
    
    current_response = initial_response
    embeddings, vectorizer = embed(current_response)
    embeddings = [embeddings]
    max_diff_response = initial_response
    max_diff = 0

    if knowledge_len > 0:
        for i in range(1, iterations):
            start_index = (i * chunk_size) % knowledge_len
            end_index = min(start_index + chunk_size, knowledge_len)
            truncated_knowledge = internet_knowledge[start_index:end_index]
            
            feedback_prompt = (
                f"Overall Context: {global_context}\n\n"
                f"Internet Knowledge (truncated): {truncated_knowledge}\n\n"
                f"Previous Outputs: {previous_output_context}\n\n"
                f"The code language is {language}"
                f"Based on this output: '{current_response}', identify any weaknesses, missing information, or potential error of the code. "
                f"Provide feedback in 7-9 bullet points, make sure the points are mid-long detailed sentences."
            )
            feedback = llm_generate(feedback_prompt, model=model)
            time.sleep(np.random.randint(4, 7))
            thinking_logs.append(feedback)
            
            revision_prompt = (
                f"Overall Context: {global_context}\n\n"
                f"Internet Knowledge (truncated): {truncated_knowledge}\n\n"
                f"Previous Outputs: {previous_output_context}\n\n"
                f"The code language is {language}"
                f"Taking into account the following feedback: '{feedback}', revise and improve this output: "
                f"'{current_response}' based on the initial prompt: '{initial_prompt}'.\n"
                "The Format should be:\n"
                "Variables : The variables of the code and one (short-mid length) sentence description of it."
                "Code:\n[Give The Finished Code of the Part] \n\n"
                "FOLLOW STRICTLY THE CODE SYNTAX SO THE CODE WORKS."
            )
            current_response = llm_generate(revision_prompt, model=model)
            time.sleep(np.random.randint(4, 7))
            all_responses.append(current_response)
            
            current_embedding = embed(current_response, existing_vectorizer=vectorizer)
            embeddings.append(current_embedding)
            similarity_diff = 1 - cosine_similarity(embeddings[0], current_embedding)[0][0]
            if similarity_diff > max_diff:
                max_diff = similarity_diff
                max_diff_response = current_response

    final_prompt = (
        f"The code language is {language}"
        f"{initial_prompt}\n\nTaking into account the following feedback:\n{''.join(thinking_logs[-4:])}\n\n"
        f"And considering the previous best response:\n{current_response}\n"
        "The Format should be:\n"
        "Variables : The variables of the code and one (short-mid length) sentence description of it."
        "Code:\n[Give The Finished Code of the Part]\n\n"
        "FOLLOW STRICTLY THE CODE SYNTAX SO THE CODE WORKS."
    )
    final_response = llm_generate(final_prompt, model=model)
    time.sleep(np.random.randint(4, 7))
    final_embedding = embed(final_response, existing_vectorizer=vectorizer)
    embeddings.append(final_embedding)
    visualize_iterative_embeddings(embeddings)
    
    return [max_diff_response], [initial_response], thinking_logs

# ---------------------------------------
# Output File Conversion Functions
# ---------------------------------------

def save_results_to_excel_in_memory(parts_list, refined_results):
    """
    Saves refined results for each TOC part to an Excel file.
    Returns the binary Excel data.
    """
    rows = []
    for idx, part in enumerate(parts_list):
        row = {
            "PART": part,
            "FINAL OUTPUT": refined_results[idx]["FINAL OUTPUT"],
            "INITIAL OUTPUT": refined_results[idx]["INITIAL OUTPUT"],
            "THINKING LOGS": "\n\n".join(refined_results[idx]["THINKING LOGS"])
        }
        rows.append(row)
    df_results = pd.DataFrame(rows)
    
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        df_results.to_excel(writer, index=False)
    excel_buffer.seek(0)
    return excel_buffer.getvalue()

def save_results_to_docx_in_memory(parts_list, refined_results, base_filename="refined_journal_results"):
    """
    Saves the parts, final output, initial output, and thinking logs into separate DOCX documents.
    Returns a dictionary of filenames to binary data.
    """
    files = {}
    
    # Save parts
    doc_part = Document()
    doc_part.add_heading("PART", level=1)
    for idx, part in enumerate(parts_list, start=1):
        doc_part.add_heading(f"Part {idx}: {part}", level=2)
    buffer_part = BytesIO()
    doc_part.save(buffer_part)
    buffer_part.seek(0)
    files[f"{base_filename}_part.docx"] = buffer_part.getvalue()
    
    # Save final outputs
    doc_final = Document()
    doc_final.add_heading("FINAL OUTPUT", level=1)
    for idx, res in enumerate(refined_results, start=1):
        doc_final.add_heading(f"Part {idx}", level=2)
        doc_final.add_paragraph(res["FINAL OUTPUT"])
    buffer_final = BytesIO()
    doc_final.save(buffer_final)
    buffer_final.seek(0)
    files[f"{base_filename}_final_output.docx"] = buffer_final.getvalue()
    
    # Save initial outputs
    doc_initial = Document()
    doc_initial.add_heading("INITIAL OUTPUT", level=1)
    for idx, res in enumerate(refined_results, start=1):
        doc_initial.add_heading(f"Part {idx}", level=2)
        doc_initial.add_paragraph(res["INITIAL OUTPUT"])
    buffer_initial = BytesIO()
    doc_initial.save(buffer_initial)
    buffer_initial.seek(0)
    files[f"{base_filename}_initial_output.docx"] = buffer_initial.getvalue()
    
    # Save thinking logs
    doc_logs = Document()
    doc_logs.add_heading("THINKING LOGS", level=1)
    for idx, res in enumerate(refined_results, start=1):
        doc_logs.add_heading(f"Part {idx}", level=2)
        doc_logs.add_paragraph("\n\n".join(res["THINKING LOGS"]))
    buffer_logs = BytesIO()
    doc_logs.save(buffer_logs)
    buffer_logs.seek(0)
    files[f"{base_filename}_thinking_logs.docx"] = buffer_logs.getvalue()
    
    return files

def save_results_to_txt_in_memory(parts_list, refined_results, base_filename="refined_journal_results"):
    """
    Saves the parts, final outputs, initial outputs, and thinking logs into TXT files.
    Returns a dictionary of filenames to binary data.
    """
    files = {}
    
    parts_txt = "\n".join(parts_list)
    files[f"{base_filename}_part.txt"] = parts_txt.encode("utf-8")
    
    final_txt = "\n".join([res["FINAL OUTPUT"] for res in refined_results])
    files[f"{base_filename}_final_output.txt"] = final_txt.encode("utf-8")
    
    initial_txt = "\n".join([res["INITIAL OUTPUT"] for res in refined_results])
    files[f"{base_filename}_initial_output.txt"] = initial_txt.encode("utf-8")
    
    logs_txt = "\n\n".join(["\n\n".join(res["THINKING LOGS"]) for res in refined_results])
    files[f"{base_filename}_thinking_logs.txt"] = logs_txt.encode("utf-8")
    
    return files

# ---------------------------------------
# Session State Reset Function
# ---------------------------------------

def reset_toc_state(session_state):
    """
    Resets TOC-related keys in the session state.
    Expects session_state to be a dictionary-like object.
    """
    for key in ["toc", "toc_locked", "toc_text_area", "proceed", "process_done"]:
        if key in session_state:
            del session_state[key]

# ============================================ STREAMLIT APP CODE ==================================================================================

if "model_configured" not in st.session_state:
    st.session_state["model"] = configure_model()
    st.session_state["model_configured"] = True

colmain1, colmain2 = st.columns([1, 12])
with colmain1:
    image = Image.open('logo.png')
    resized_image = image.resize((80, 80))
    st.image(resized_image)
with colmain2:
    st.markdown(
        """
        <span style='color:white; font-size:36px; font-weight:bold; margin-top: 15px; display: block;'>n_secondbrain-codegen</span>
        """, 
        unsafe_allow_html=True
    )

# ----------------------------------------------------------------------------- 
# Layout Setup: Three Columns ([3,2,2]) 
# -----------------------------------------------------------------------------
col1, col2, col3 = st.columns([3, 2, 2])

# Column 1: Prompt Input and Length Info
with col1.expander("", expanded=True):
    st.subheader("Select Code Language")
    language = st.selectbox(
        "",
        ("Python", "JavaScript", "HTML", 
         "CSS", "C", "C++", 
         "C#", "Rust", "R"),
    )
    st.subheader("Enter Main Prompt")
    prompt_input = st.text_area("Main Prompt", height=410)
    prompt_length_container = st.container()
    if prompt_input:
        prompt_len = len(prompt_input)
        prompt_length_container.info(f"Prompt length: {prompt_len} / 10000")
        if prompt_len > 10000:
            prompt_length_container.error("Prompt exceeds limit")
    else:
        prompt_length_container.info("Prompt length: 0 / 10000")

# Column 2: TOC and Refinement Settings
with col2.expander("", expanded=True):
    st.subheader("Table of Contents Settings")
    st.write("Select how many parts the task should be divided into.")
    toc_parts = st.slider("", min_value=1, max_value=40, value=5)
    
    st.subheader("Refinement Iterations")
    st.write("Select how many refinement iterations (thinking sessions) to perform for each part.")
    n_iterations = st.slider("", min_value=1, max_value=7, value=4)

    st.subheader("Internet Search")
    use_internet = st.checkbox("Use Internet Search as Context")
    if use_internet:
        after_date = st.date_input("Search After", value=datetime.date(2020, 1, 1))
        before_date = st.date_input("Search Before", value=datetime.date(2029, 1, 1))

# Column 3: Output File Settings & Initial Generate Button
with col3.expander("", expanded=True):
    st.subheader("Select Types of File")
    st.write("Choose the file formats for the final outputs (docx, excel, txt).")
    doc_types = st.multiselect(
        "Document Types",
        options=["docx", "excel", "txt"],
        default=["docx", "excel", "txt"]
    )
    if prompt_input:
        est_time = toc_parts * n_iterations * 0.8
        st.warning(f"Estimated processing time: ~{est_time:.1f} minutes.")
    if prompt_input and "toc_locked" not in st.session_state:
        generate_toc_button = st.button("Generate TOC", key="generate_toc")

# ----------------------------------------------------------------------------- 
# --- TOC Generation & Display --- 
# -----------------------------------------------------------------------------
if prompt_input and "toc_locked" not in st.session_state and generate_toc_button:
    st.session_state["toc"] = table_of_contents(prompt_input, n_parts=toc_parts, model=st.session_state["model"])
    st.success("Table of Contents generated!")

if prompt_input and ("toc" in st.session_state or "toc_locked" in st.session_state):
    with st.expander("", expanded=True):
        st.subheader("Generated Table of Contents")
        st.write("This Table of Contents was generated to give you an idea:")
        if "toc_locked" in st.session_state:
            st.write(st.session_state["toc_locked"])
        else:
            st.write(st.session_state["toc"])
            if st.button("Re-Generate TOC", key="regen_toc"):
                st.session_state["toc"] = table_of_contents(prompt_input, n_parts=toc_parts, model=st.session_state["model"])
                st.success("Table of Contents regenerated!")
    
    if "toc_locked" not in st.session_state:
        with st.expander("", expanded=True):
            st.subheader("Table of Contents Customization")
            if "toc_text_area" not in st.session_state:
                st.session_state["toc_text_area"] = st.session_state["toc"]
            locked_toc = st.text_area("Edit Your Table of Contents (or copy and paste it here) (note: MUST HAVE NUMBER BEFORE PART (ex : 1. Part 1 ; 2. Part 2)):", 
                                            value=st.session_state["toc_text_area"],
                                            height=200, key="toc_text_area")
            if st.button("Process", key="process_toc"):
                st.session_state["toc_locked"] = locked_toc
                st.session_state["proceed"] = True

# ----------------------------------------------------------------------------- 
# Proceed with Sequential Process (only if locked TOC is set and process not done)
# -----------------------------------------------------------------------------
if (prompt_input 
    and st.session_state.get("toc_locked") 
    and st.session_state.get("proceed") 
    and not st.session_state.get("process_done")):
    
    st.info("Processing started... This may take a while. Do not press any buttons and sit still!")
    
    model = st.session_state["model"]
    # Step 1: Extract parts from the locked TOC
    parts_list = extract_parts(st.session_state["toc_locked"])
    
    with st.expander("", expanded=True):
        st.subheader("Progress Log")
        progress_container = st.empty()
        debug_log = ""
        
        # Step 2: Generate Detailed Prompts for each TOC part
        debug_log = "Generating Detailed Prompts...\n"
        progress_container.info(debug_log)
        global_context = f"Overall Table of Contents: {st.session_state['toc_locked']}\n\n"
        content_prompts = generate_prompts_for_parts(parts_list, model=model, global_context=global_context)
        debug_log = "Detailed Prompts generated.\n"
        progress_container.info(debug_log)
        
        # Step 3: Internet Search (if enabled) for each TOC part
        internet_knowledge = {}
        excel_dataframes = {}
        queries_by_part = {} # Initialize the variable
        if use_internet:
            queries_by_part = make_prompt_queries(prompt_input, n_parts=len(parts_list), model=model) #get queries
            
            for idx, part in enumerate(parts_list, start=1):
                debug_log = f"\nStarting data collection for TOC part {idx}\n"
                progress_container.info(debug_log)
                part_content = ""
                part_data = []
                
                # For each of the queries for this part
                for query in queries_by_part.get(f"Part {idx}", []):
                    # Append date filters to the query
                    formatted_query = f"{query} after:{after_date} before:{before_date}"
                    debug_log = f"Part {idx}: Searching with query: {formatted_query}\n"
                    progress_container.info(debug_log)
                    
                    news_data = getNewsData(formatted_query)
                    if not news_data:
                        debug_log = f"Warning: No news data returned for query: {formatted_query}\n"
                        progress_container.info(debug_log)
                    
                    links = [item["link"] for item in news_data]
                    articles_info = extract_article_info(links)
                    
                    for article in articles_info:
                        if article:  # Check if article is not None (i.e., content length was sufficient)
                            part_content += article["content"] + "\n"
                            part_data.append({
                                "query": formatted_query,
                                "title": article["title"],
                                "summary": article["summary"],
                                "content": article["content"],
                                "link": article["link"]
                            })
                    debug_log = f"After query '{query}', downloaded articles count for Part {idx}: {len(part_data)}\n"
                    progress_container.info(debug_log)
                    time.sleep(1)

                    # Store the internet knowledge and dataframe for this part
                    internet_knowledge[part] = part_content
                    df = pd.DataFrame(part_data)
                    excel_dataframes[f"internetsearch_part{idx}"] = df
                debug_log = f"Final article count for TOC Part {idx}: {len(df)}\n"
                progress_container.info(debug_log)
        
        # Step 4: Iterative Refinement for each TOC part
        refined_results = []
        total_parts = len(content_prompts)
        previous_part_outputs = [] # Initialize list to store previous outputs

        for idx, cp in enumerate(content_prompts, start=0):
            debug_log = f"Processing iterative refinement for Part {idx+1}/{total_parts}...\n"
            progress_container.info(debug_log)
            
            # Get the internet knowledge (content) for the corresponding part; if not available, pass an empty string
            current_internet_knowledge = internet_knowledge.get(parts_list[idx], "")
            
            max_response, init_response, thinking_logs = iterative_refinement(
                cp, 
                internet_knowledge=current_internet_knowledge, 
                previous_outputs=previous_part_outputs, # Pass previous outputs
                iterations=n_iterations, 
                global_context=global_context, 
                model=model
            )
            refined_results.append({
                "FINAL OUTPUT": max_response[0],
                "INITIAL OUTPUT": init_response[0],
                "THINKING LOGS": thinking_logs
            })

            previous_part_outputs.append(max_response[0]) #store output for next part

            debug_log = f"Finished refinement for Part {idx+1}\n"
            progress_container.info(debug_log)
            time.sleep(1)
        progress_container.empty()
        st.success("Refinement complete!")
    
    # Step 5: Generate in-memory output files based on selected types
    files = {}
    if "excel" in doc_types:
        # Save each scraped dataframe as its own Excel file
        for df_name, df in excel_dataframes.items():
            excel_buffer = BytesIO()
            with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
                df.to_excel(writer, index=False)
            excel_buffer.seek(0)
            files[f"{df_name}.xlsx"] = excel_buffer.getvalue()
        
        # Also save the refined results into a consolidated Excel file
        df_results = pd.DataFrame([{
            "PART": part,
            "FINAL OUTPUT": refined_results[idx]["FINAL OUTPUT"],
            "INITIAL OUTPUT": refined_results[idx]["INITIAL OUTPUT"],
            "THINKING LOGS": "\n\n".join(refined_results[idx]["THINKING LOGS"])
        } for idx, part in enumerate(parts_list)])
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
            df_results.to_excel(writer, index=False)
        excel_buffer.seek(0)
        files["refined_results.xlsx"] = excel_buffer.getvalue()
    
    if "docx" in doc_types:
        docx_files = save_results_to_docx_in_memory(parts_list, refined_results, base_filename="my_results")
        files.update(docx_files)
    if "txt" in doc_types:
        txt_files = save_results_to_txt_in_memory(parts_list, refined_results, base_filename="my_results")
        files.update(txt_files)
    
    # Step 6: Zip all in-memory files (without writing to disk)
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        for filename, filedata in files.items():
            zipf.writestr(filename, filedata)
    zip_buffer.seek(0)
    st.session_state["download_zip"] = zip_buffer
    st.success("All outputs generated and zipped successfully!")
    st.session_state["ready_to_download"] = True
    st.session_state["process_done"] = True

    if use_internet:
        with st.expander("Generated Queries for Each TOC Part (Debug Info)",expanded=True):
            for part, queries in queries_by_part.items():
                st.write(f"**{part}:**")
                for query in queries:
                    st.write(f"- {query}")

# ----------------------------------------------------------------------------- 
# Download Container (Single Download Button) 
# -----------------------------------------------------------------------------
if st.session_state.get("ready_to_download", False):
    st.download_button(
        label="Download All Outputs (ZIP)",
        data=st.session_state["download_zip"],
        file_name="refined_document_outputs.zip",
        mime="application/zip",
        on_click=lambda: reset_toc_state(st.session_state)
    )
